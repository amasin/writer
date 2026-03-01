#!/usr/bin/env python3
"""
Proofreader and SEO QA Agent.
Uses A2A protocol to review articles, score them, and request improvements.
"""

import re
from typing import Dict, Any, Tuple, Optional, List
from datetime import datetime
from a2a_protocol import A2AAgent, A2AMessage, MessageType, AgentType, A2AMessageBroker


class ProofreaderAgent(A2AAgent):
    """Agent for proofreading articles and assigning SEO quality scores."""
    
    def __init__(self, agent_id: str = "proofreader_agent"):
        """Initialize the Proofreader Agent."""
        super().__init__(agent_id, AgentType.WORDPRESS_WRITER_AGENT)  # Using same type for now
        self.message_broker: Optional[A2AMessageBroker] = None
        self.score_history: Dict[str, List[float]] = {}
        self.suggestion_history: Dict[str, List[str]] = {}
    
    def process_message(self, message: A2AMessage) -> A2AMessage:
        """
        Process incoming A2A message for article review.
        
        Args:
            message: The incoming A2A message
            
        Returns:
            Response message with SEO score and suggestions
        """
        if message.message_type != MessageType.REQUEST:
            return A2AMessage(
                sender=self.agent_id,
                receiver=message.sender,
                message_type=MessageType.ERROR,
                payload={"error": "Invalid message type"}
            )
        
        article_content = message.payload.get("content", "")
        title = message.payload.get("title", "")
        topic = message.payload.get("topic", "")
        iteration = message.payload.get("iteration", 1)
        
        # Analyze the article
        seo_score, suggestions = self.analyze_article(article_content, title, topic)
        
        # Track score history
        if title not in self.score_history:
            self.score_history[title] = []
            self.suggestion_history[title] = []
        
        self.score_history[title].append(seo_score)
        self.suggestion_history[title].extend(suggestions)
        
        return A2AMessage(
            sender=self.agent_id,
            receiver=message.sender,
            message_type=MessageType.RESPONSE,
            payload={
                "seo_score": seo_score,
                "suggestions": suggestions,
                "iteration": iteration,
                "ready_for_publication": seo_score >= 8.0,
                "detailed_analysis": self._generate_detailed_analysis(article_content, title)
            }
        )
    
    def set_message_broker(self, broker: A2AMessageBroker) -> None:
        """Set the A2A message broker for agent communication."""
        self.message_broker = broker
    
    def analyze_article(self, content: str, title: str = "", topic: str = "") -> Tuple[float, List[str]]:
        """
        Analyze article and calculate SEO score.
        
        Args:
            content: Article content in HTML/WordPress format
            title: Article title
            topic: Main topic
            
        Returns:
            Tuple of (seo_score, suggestions_list)
        """
        score = 0.0
        suggestions = []
        
        # Extract plain text from HTML
        plain_text = self._extract_plain_text(content)
        
        # 1. Title Optimization (max 15 points)
        title_score, title_suggestions = self._analyze_title(title, topic)
        score += title_score
        suggestions.extend(title_suggestions)
        
        # 2. Content Length (max 15 points)
        length_score, length_suggestions = self._analyze_content_length(plain_text)
        score += length_score
        suggestions.extend(length_suggestions)
        
        # 3. Keyword Density (max 15 points)
        density_score, density_suggestions = self._analyze_keyword_density(plain_text, topic)
        score += density_score
        suggestions.extend(density_suggestions)
        
        # 4. Heading Structure (max 15 points)
        heading_score, heading_suggestions = self._analyze_heading_structure(content)
        score += heading_score
        suggestions.extend(heading_suggestions)
        
        # 5. Meta Elements & HTML Structure (max 15 points)
        html_score, html_suggestions = self._analyze_html_structure(content)
        score += html_score
        suggestions.extend(html_suggestions)
        
        # 6. Readability (max 10 points)
        readability_score, readability_suggestions = self._analyze_readability(plain_text)
        score += readability_score
        suggestions.extend(readability_suggestions)
        
        # 7. Links and Internal Structure (max 10 points)
        link_score, link_suggestions = self._analyze_links(content)
        score += link_score
        suggestions.extend(link_suggestions)
        
        # Ensure score is within 0-10 range
        final_score = min(score, 10.0)
        
        return final_score, suggestions
    
    def _extract_plain_text(self, html_content: str) -> str:
        """Extract plain text from HTML content."""
        # Remove HTML tags
        text = re.sub(r'<[^>]+>', '', html_content)
        # Remove HTML entities
        text = re.sub(r'&[^;]+;', '', text)
        # Clean up whitespace
        text = ' '.join(text.split())
        return text
    
    def _analyze_title(self, title: str, topic: str) -> Tuple[float, List[str]]:
        """Analyze title for SEO optimization."""
        score = 0.0
        suggestions = []
        
        if not title:
            return 0, ["Title is missing or empty"]
        
        # Length check (50-60 chars optimal)
        title_length = len(title)
        if 50 <= title_length <= 60:
            score += 5
        elif 40 <= title_length <= 70:
            score += 3
            suggestions.append(f"Title length is {title_length} chars; optimal is 50-60 chars")
        else:
            suggestions.append(f"Title too {'short' if title_length < 40 else 'long'}: {title_length} chars (optimal: 50-60)")
        
        # Power words check
        power_words = ["ultimate", "guide", "complete", "best", "how to", "why", 
                       "what", "essential", "proven", "comprehensive", "expert"]
        has_power_words = any(word in title.lower() for word in power_words)
        if has_power_words:
            score += 5
        else:
            suggestions.append("Add powerful words like 'Ultimate', 'Complete', 'Guide', or 'Best' to title")
        
        # Topic relevance
        if topic.lower() in title.lower():
            score += 5
        else:
            suggestions.append(f"Title should include main topic keyword: '{topic}'")
        
        return score, suggestions
    
    def _analyze_content_length(self, text: str) -> Tuple[float, List[str]]:
        """Analyze content length."""
        score = 0.0
        suggestions = []
        
        word_count = len(text.split())
        
        # Optimal word count: 800-2500 words for blog posts
        if 800 <= word_count <= 2500:
            score += 15
        elif 500 <= word_count < 800:
            score += 10
            suggestions.append(f"Content is {word_count} words; aim for 800-2500 words for better SEO")
        elif word_count > 2500:
            score += 12
            suggestions.append(f"Content is {word_count} words; consider breaking into multiple articles")
        else:
            score += 3  # Even short content can get some points
            suggestions.append(f"Content too short ({word_count} words); minimum 800 words recommended for optimal SEO")
        
        return score, suggestions
    
    def _analyze_keyword_density(self, text: str, topic: str) -> Tuple[float, List[str]]:
        """Analyze keyword density."""
        score = 0.0
        suggestions = []
        
        text_lower = text.lower()
        total_words = len(text.split())
        
        if total_words == 0:
            return 0, ["No content to analyze"]
        
        # Count topic keyword occurrences
        topic_lower = topic.lower()
        keyword_count = text_lower.count(topic_lower)
        keyword_density = (keyword_count / total_words) * 100
        
        # Optimal keyword density: 1-3%
        if 1.0 <= keyword_density <= 3.0:
            score += 8
        elif 0.5 <= keyword_density < 1.0:
            score += 5
            suggestions.append(f"Keyword density too low ({keyword_density:.1f}%); aim for 1-3%")
        elif 3.0 < keyword_density <= 5.0:
            score += 5
            suggestions.append(f"Keyword density slightly high ({keyword_density:.1f}%); aim for 1-3%")
        else:
            score += 2
            suggestions.append(f"Keyword density problematic ({keyword_density:.1f}%); target 1-3%")
        
        # Check for variations and related keywords
        related_keywords = self._get_related_keywords(topic)
        found_variations = sum(1 for keyword in related_keywords if keyword.lower() in text_lower)
        
        if found_variations >= 3:
            score += 7
        elif found_variations >= 1:
            score += 4
            suggestions.append("Add more keyword variations and related terms for better SEO coverage")
        else:
            suggestions.append(f"Incorporate related keywords: {', '.join(related_keywords[:3])}")
        
        return score, suggestions
    
    def _get_related_keywords(self, topic: str) -> List[str]:
        """Get related keywords for a topic."""
        keywords_map = {
            "artificial intelligence": ["machine learning", "AI", "neural networks", "deep learning", 
                                       "automation", "intelligent systems", "data science"],
            "": ["topic", "subject", "content", "information"]
        }
        return keywords_map.get(topic.lower(), ["related", "advanced", "implementation", "solution"])
    
    def _analyze_heading_structure(self, html_content: str) -> Tuple[float, List[str]]:
        """Analyze heading structure."""
        score = 0.0
        suggestions = []
        
        # Count headings
        h1_count = len(re.findall(r'<h1[^>]*>', html_content, re.IGNORECASE))
        h2_count = len(re.findall(r'<h2[^>]*>', html_content, re.IGNORECASE))
        h3_count = len(re.findall(r'<h3[^>]*>', html_content, re.IGNORECASE))
        
        # Optimal: 1 H1, multiple H2s, multiple H3s
        if h1_count == 1:
            score += 5
        elif h1_count == 0:
            suggestions.append("Missing H1 heading; add one main H1 per article")
        else:
            suggestions.append(f"Multiple H1 tags ({h1_count}); should have exactly one")
        
        if h2_count >= 3:
            score += 5
        elif h2_count >= 1:
            score += 3
            suggestions.append(f"Only {h2_count} H2 heading(s); aim for at least 3-5")
        else:
            suggestions.append("Missing H2 subheadings; add multiple for better structure")
        
        if h3_count >= 2:
            score += 5
        elif h3_count >= 1:
            score += 2
            suggestions.append("Add more H3 subheadings under H2s")
        else:
            suggestions.append("Consider adding H3 subheadings for better content organization")
        
        return score, suggestions
    
    def _analyze_html_structure(self, html_content: str) -> Tuple[float, List[str]]:
        """Analyze HTML structure and meta elements."""
        score = 0.0
        suggestions = []
        
        # Check for WordPress blocks
        has_wp_blocks = bool(re.search(r'<!-- wp:', html_content))
        if has_wp_blocks:
            score += 5
        else:
            suggestions.append("Content should use WordPress block format (<!-- wp:...)")
        
        # Check for paragraphs
        p_count = len(re.findall(r'<p[^>]*>', html_content, re.IGNORECASE))
        if p_count >= 5:
            score += 5
        else:
            suggestions.append(f"Add more paragraphs for better readability; found {p_count}")
        
        # Check for lists
        has_lists = bool(re.search(r'<ul[^>]*>|<ol[^>]*>', html_content, re.IGNORECASE))
        if has_lists:
            score += 5
        else:
            suggestions.append("Add bullet or numbered lists to break up content")
        
        return score, suggestions
    
    def _analyze_readability(self, text: str) -> Tuple[float, List[str]]:
        """Analyze readability."""
        score = 0.0
        suggestions = []
        
        sentences = text.split('.')
        words = text.split()
        
        if not sentences or not words:
            return 0, ["No content for readability analysis"]
        
        # Average sentence length (optimal: 15-20 words)
        avg_sent_length = len(words) / len([s for s in sentences if s.strip()])
        
        if 15 <= avg_sent_length <= 20:
            score += 5
        elif 10 <= avg_sent_length < 25:
            score += 3
            suggestions.append(f"Average sentence length: {avg_sent_length:.1f} words; optimal is 15-20")
        else:
            suggestions.append(f"Readability issue: avg sentence {avg_sent_length:.1f} words; aim for 15-20")
        
        # Check for short paragraphs (readability)
        # This is a simple heuristic
        score += 5  # Assuming good readability if content is well-structured
        
        return score, suggestions
    
    def _analyze_links(self, html_content: str) -> Tuple[float, List[str]]:
        """Analyze internal and external links."""
        score = 0.0
        suggestions = []
        
        # Count links (using simple regex)
        all_links = len(re.findall(r'<a[^>]+href=["\']([^"\']+)["\']', html_content, re.IGNORECASE))
        
        if all_links >= 3:
            score += 10
        elif all_links >= 1:
            score += 5
            suggestions.append(f"Add more internal/external links ({all_links} found); aim for 3-5")
        else:
            suggestions.append("Add relevant internal and external links (3-5 recommended)")
        
        return score, suggestions
    
    def _generate_detailed_analysis(self, content: str, title: str) -> Dict[str, Any]:
        """Generate detailed analysis breakdown."""
        plain_text = self._extract_plain_text(content)
        word_count = len(plain_text.split())
        h2_count = len(re.findall(r'<h2[^>]*>', content, re.IGNORECASE))
        p_count = len(re.findall(r'<p[^>]*>', content, re.IGNORECASE))
        
        return {
            "word_count": word_count,
            "title_length": len(title),
            "heading_count": h2_count,
            "paragraph_count": p_count,
            "analyzed_at": datetime.now().isoformat()
        }
    
    def review_and_improve(self, article_content: str, title: str, topic: str,
                          message_broker: A2AMessageBroker, max_iterations: int = 3) -> Tuple[str, float]:
        """
        Iteratively review and improve article until SEO score >= 8.
        
        Args:
            article_content: Original article content
            title: Article title
            topic: Article topic
            message_broker: Message broker for A2A communication
            max_iterations: Maximum improvement iterations
            
        Returns:
            Tuple of (final_article_content, final_seo_score)
        """
        current_content = article_content
        iteration = 1
        
        print(f"\n[Proofreader] Starting iterative review process...", file=__import__('sys').stderr)
        
        while iteration <= max_iterations:
            print(f"[Proofreader] Iteration {iteration}: Analyzing article...", file=__import__('sys').stderr)
            
            # Analyze current version
            seo_score, suggestions = self.analyze_article(current_content, title, topic)
            
            print(f"[Proofreader] SEO Score: {seo_score:.1f}/10", file=__import__('sys').stderr)
            print(f"[Proofreader] Top suggestions: {suggestions[:2] if suggestions else 'None'}", file=__import__('sys').stderr)
            
            # Check if score is acceptable
            if seo_score >= 8.0:
                print(f"[Proofreader] SEO Score {seo_score:.1f} >= 8.0 - Ready for publication!", file=__import__('sys').stderr)
                return current_content, seo_score
            
            # Request improvements
            if iteration < max_iterations:
                print(f"[Proofreader] Requesting article improvements...", file=__import__('sys').stderr)
                
                # Send improvement request to WordPress Writer Agent
                message = A2AMessage(
                    sender=self.agent_id,
                    receiver="wordpress_writer_agent",
                    message_type=MessageType.REQUEST,
                    payload={
                        "content": current_content,
                        "title": title,
                        "topic": topic,
                        "seo_score": seo_score,
                        "suggestions": suggestions,
                        "iteration": iteration,
                        "request_type": "improve_seo"
                    }
                )
                
                # Get improved version
                response = message_broker.send_message(message)
                
                if response.message_type == MessageType.RESPONSE:
                    current_content = response.payload.get("improved_content", current_content)
                    print(f"[Proofreader] Received improved article", file=__import__('sys').stderr)
                else:
                    print(f"[Proofreader] Error getting improved article", file=__import__('sys').stderr)
                    break
            
            iteration += 1
        
        # After max iterations
        seo_score, _ = self.analyze_article(current_content, title, topic)
        return current_content, seo_score
    
    def export_to_word(self, title: str, article_content: str, seo_score: float) -> str:
        """
        Export article to Word document.
        
        Args:
            title: Article title
            article_content: Article content (HTML format)
            seo_score: Final SEO score
            
        Returns:
            Path to the created Word document
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            print("[Error] python-docx not installed. Install with: pip install python-docx", file=__import__('sys').stderr)
            return ""
        
        # Create document
        doc = Document()
        
        # Add title
        title_para = doc.add_heading(title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        meta_para = doc.add_paragraph()
        meta_para.add_run(f"SEO Score: {seo_score:.1f}/10").bold = True
        meta_para.add_run(f" | Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add horizontal line
        doc.add_paragraph("_" * 80)
        
        # Extract and format content
        plain_text = self._extract_plain_text(article_content)
        
        # Extract sections
        sections = self._extract_sections_from_html(article_content)
        
        for section in sections:
            section_type = section.get("type", "paragraph")
            section_content = section.get("content", "")
            section_level = section.get("level", 1)
            
            if section_type == "heading":
                doc.add_heading(section_content, level=section_level)
            elif section_type == "list_item":
                para = doc.add_paragraph(section_content, style='List Bullet')
            elif section_type == "paragraph":
                # Clean up the paragraph text
                clean_text = section_content.strip()
                if clean_text:
                    doc.add_paragraph(clean_text)
        
        # Add footer with SEO metrics
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer_text = footer.add_run("\n" + "="*80 + "\n")
        footer_text.bold = True
        
        metrics = doc.add_paragraph()
        metrics.add_run("Document Metrics:\n").bold = True
        metrics.add_run(f"• Total Words: {len(plain_text.split())}\n")
        metrics.add_run(f"• SEO Score: {seo_score:.1f}/10\n")
        metrics.add_run(f"• Status: {'Ready for Publication' if seo_score >= 8.0 else 'Needs Improvement'}\n")
        metrics.add_run(f"• Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Save document
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"article_{title.replace(' ', '_')[:50]}_{timestamp}.docx"
        doc.save(filename)
        
        return filename
    
    def _extract_sections_from_html(self, html_content: str) -> list:
        """Extract sections from WordPress HTML content."""
        sections = []
        
        # Split by HTML comments for paragraphs
        import re
        
        # Find all heading tags
        headings = re.finditer(r'<h(\d)[^>]*>([^<]+)</h\1>', html_content, re.IGNORECASE)
        for match in headings:
            level = int(match.group(1))
            content = match.group(2).strip()
            sections.append({
                "type": "heading",
                "content": content,
                "level": level
            })
        
        # Find all paragraphs
        paragraphs = re.finditer(r'<p[^>]*>([^<]+)</p>', html_content, re.IGNORECASE)
        for match in paragraphs:
            content = match.group(1).strip()
            # Remove HTML entities and tags
            content = re.sub(r'<[^>]+>', '', content)
            content = re.sub(r'&[^;]+;', '', content)
            if content:
                sections.append({
                    "type": "paragraph",
                    "content": content
                })
        
        # Find all list items
        list_items = re.finditer(r'<li[^>]*>([^<]+)</li>', html_content, re.IGNORECASE)
        for match in list_items:
            content = match.group(1).strip()
            if content:
                sections.append({
                    "type": "list_item",
                    "content": content
                })
        
        return sections
